import part_1 as part1
import part_2 as part2
import part_3 as part3

import time

# from tensorflow.keras.datasets import mnist
from sklearn.model_selection import train_test_split
import tensorflow
import numpy as np

def get_dataset():
    (X_train_full, y_train_full), (X_test, y_test) = tensorflow.keras.datasets.mnist.load_data()
    # scale image values to [0, 1]
    X_train_full = X_train_full.astype('float32') / 255
    X_test = X_test.astype('float32') / 255

    X_train, X_val, y_train, y_val = train_test_split(X_train_full, y_train_full, test_size=0.2, random_state=42)

    # Reshape from 2d to 1d
    X_train = X_train.reshape(-1, 28 * 28).T
    X_val = X_val.reshape(-1, 28 * 28).T
    X_test = X_test.reshape(-1, 28 * 28).T

    # create a one hot encoded matrix and transpose it to the matching shape of the softmax output
    Y_train = tensorflow.keras.utils.to_categorical(y_train).T  # shape: (10, num_train_examples)
    Y_val = tensorflow.keras.utils.to_categorical(y_val).T
    Y_test = tensorflow.keras.utils.to_categorical(y_test).T

    return X_train, Y_train, X_val, Y_val, X_test, Y_test

# def train_model_section_4(X_train, Y_train, X_val, Y_val, X_test, Y_test, layers, learning_rate, batch_size, epochs):
#     max_val_accuracy = 0
#     no_improve_count = 0
#     num_steps_without_improvement = 100
#     small_improvement = 1e-4
#     costs = []
#     num_of_iterations = X_train.shape[1] // batch_size
#     parameters = None
#     for epoch in range(epochs):
#         print(f"Epoch {epoch + 1}/{epochs}")
#         parameters, epoch_costs = part3.l_layer_model(X_train, Y_train, layers, learning_rate, num_of_iterations, batch_size, use_batchnorm=False, use_l2=False, parameters=parameters)
#         costs.extend((epoch,epoch_costs))
#
#         val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=False)
#         print(f"Validation accuracy after epoch {epoch + 1}: {val_accuracy:.2f}%")
#         if val_accuracy > max_val_accuracy + small_improvement:
#             max_val_accuracy = val_accuracy
#             no_improve_count = 0
#         else:
#             no_improve_count += 1
#
#         if no_improve_count > num_steps_without_improvement:
#             print("Early stop no improvement on the validation")
#             break
#
#     print("Training finished.")
#     train_accuracy = part3.predict(X_train, Y_train, parameters)
#     val_accuracy = part3.predict(X_val, Y_val, parameters)
#     test_accuracy = part3.predict(X_test, Y_test, parameters)
#
#     print(f"Train accuracy: {train_accuracy:.2f}%")
#     print(f"Validation accuracy: {val_accuracy:.2f}%")
#     print(f"Test accuracy: {test_accuracy:.2f}%")
#
#     print("###############################")
#     print("Weights Norms Per Layer section 4:")
#     for l in range(1, len(layers)):
#         W = parameters[f'W{l}']
#         print(f"||W{l}||_F = {np.linalg.norm(W):.4f}")
#     return parameters, costs
#
#     # max_val_accuracy = 0
#     # no_improve_count = 0
#     # num_steps_without_improvement = 100
#     # small_improvement = 1e-4
#     # step_counter = 0
#     # costs = []
#     #
#     # parameters = None  # Pass None to l_layer_model so it will init the parameters
#     # steps_in_epoch = X_train.shape[1] // batch_size
#     #
#     # for epoch in range(epochs):
#     #     print(f"Epoch {epoch + 1}/{epochs}")
#     #
#     #     # Run l_layer_model for one epoch
#     #     parameters, epoch_costs = part3.l_layer_model(X_train, Y_train, layers,
#     #                                                    learning_rate,
#     #                                                    steps_in_epoch,
#     #                                                    batch_size,
#     #                                                    False,
#     #                                                    use_l2=False)
#     #
#     #     # Collect costs
#     #     costs.extend(epoch_costs)
#     #
#     #     # Every epoch evaluate on validation set
#     #     val_accuracy = part3.predict(X_val, Y_val, parameters)
#     #
#     #     if val_accuracy > max_val_accuracy + small_improvement:
#     #         max_val_accuracy = val_accuracy
#     #         no_improve_count = 0
#     #     else:
#     #         no_improve_count += 1
#     #
#     #     print(f"Validation accuracy after epoch {epoch + 1}: {val_accuracy:.2f}%")
#     #
#     #     if no_improve_count > num_steps_without_improvement:
#     #         print("Early stop no improvement on the validation")
#     #         break
#     #
#     # print("Training finished.")
#     #
#     # train_accuracy = part3.predict(X_train, Y_train, parameters)
#     # val_accuracy = part3.predict(X_val, Y_val, parameters)
#     # test_accuracy = part3.predict(X_test, Y_test, parameters)
#     #
#     # print(f"Train accuracy: {train_accuracy:.2f}%")
#     # print(f"Validation accuracy: {val_accuracy:.2f}%")
#     # print(f"Test accuracy: {test_accuracy:.2f}%")
#     #
#     # return parameters, costs
# ###############################################################################
#     # max_val_accuracy = 0
#     # no_improve_count = 0
#     # num_steps_without_improvement = 100
#     # step_counter = 0
#     # small_improvement = 1e-4
#     # parameters = part_1.initialize_parameters(layers)
#     # steps_in_epoch = X_train.shape[1] // batch_size
#     # costs = []
#     # for epoch in range(epochs):
#     #     print(f"Epoch {epoch + 1}/{epochs}")
#     #
#     #     for step in range(steps_in_epoch):
#     #         start = step * batch_size
#     #         end = start + batch_size
#     #         X_batch = X_train[:, start:end]
#     #         Y_batch = Y_train[:, start:end]
#     #
#     #         # Forward propagation
#     #         AL, caches = part_1.l_model_forward(X_batch, parameters, use_batchnorm=False)
#     #
#     #         # Compute cost
#     #         cost = part_1.compute_cost(AL, Y_batch)
#     #
#     #         # Backward propagation
#     #         grads = part_2.l_model_backward(AL, Y_batch, caches)
#     #
#     #         # Update parameters
#     #         parameters = part_2.update_parameters(parameters, grads, learning_rate)
#     #
#     #         step_counter += 1
#     #
#     #         if step_counter % 100 == 0:
#     #             print(f"Cost after iteration {step_counter}: {cost:.4f}")
#     #             costs.append(cost)
#     #
#     #         if step_counter % 10 == 0: # every 10 mini batches evaluate the model on the validation set
#     #             val_accuracy = part_3.predict(X_val, Y_val, parameters)
#     #
#     #             if val_accuracy > max_val_accuracy + small_improvement:
#     #                 max_val_accuracy = val_accuracy
#     #                 no_improve_count = 0
#     #             else:
#     #                 no_improve_count += 1
#     #
#     #             # print(f"Step {step_counter} | Cost: {cost:.4f} | Val Acc: {val_accuracy:.2f}%")
#     #             if no_improve_count > num_steps_without_improvement:
#     #                 print("Early stop no improvement on the validation")
#     #                 break
#     #     if no_improve_count > num_steps_without_improvement:
#     #         break
#     # print("Training finished.")
#     # train_accuracy = part_3.predict(X_train, Y_train, parameters)
#     # val_accuracy = part_3.predict(X_val, Y_val, parameters)
#     # test_accuracy = part_3.predict(X_test, Y_test, parameters)
#     #
#     # print(f"Train accuracy: {train_accuracy:.2f}%")
#     # print(f"Validation accuracy: {val_accuracy:.2f}%")
#     # print(f"Test accuracy: {test_accuracy:.2f}%")
#
#
# # Press the green button in the gutter to run the script.
# def train_model_section_5(X_train, Y_train, X_val, Y_val, X_test, Y_test, layers, learning_rate, batch_size, epochs):
#     max_val_accuracy = 0
#     no_improve_count = 0
#     num_steps_without_improvement = 100
#     small_improvement = 1e-4
#     costs = []
#     num_of_iterations = X_train.shape[1] // batch_size
#     parameters = None
#     for epoch in range(epochs):
#         print(f"Epoch {epoch + 1}/{epochs}")
#         parameters, epoch_costs = part3.l_layer_model(X_train, Y_train, layers, learning_rate, num_of_iterations, batch_size, use_batchnorm=True, use_l2=False, parameters=parameters)
#         costs.extend((epoch,epoch_costs))
#
#         val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=True)
#         print(f"Validation accuracy after epoch {epoch + 1}: {val_accuracy:.2f}%")
#         if val_accuracy > max_val_accuracy + small_improvement:
#             max_val_accuracy = val_accuracy
#             no_improve_count = 0
#         else:
#             no_improve_count += 1
#
#         if no_improve_count > num_steps_without_improvement:
#             print("Early stop no improvement on the validation")
#             break
#
#     print("Training finished.")
#     train_accuracy = part3.predict(X_train, Y_train, parameters, use_batchnorm=True)
#     val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=True)
#     test_accuracy = part3.predict(X_test, Y_test, parameters, use_batchnorm=True)
#
#     print(f"Train accuracy: {train_accuracy:.2f}%")
#     print(f"Validation accuracy: {val_accuracy:.2f}%")
#     print(f"Test accuracy: {test_accuracy:.2f}%")
#
#     return parameters, costs
#
#
#     # max_val_accuracy = 0
#     # no_improve_count = 0
#     # num_steps_without_improvement = 100
#     # step_counter = 0
#     # small_improvement = 1e-4
#     # parameters = part1.initialize_parameters(layers)
#     # steps_in_epoch = X_train.shape[1] // batch_size
#     # costs = []
#     # for epoch in range(epochs):
#     #     print(f"Epoch {epoch + 1}/{epochs}")
#     #
#     #     for step in range(steps_in_epoch):
#     #         start = step * batch_size
#     #         end = start + batch_size
#     #         X_batch = X_train[:, start:end]
#     #         Y_batch = Y_train[:, start:end]
#     #
#     #         # Forward propagation
#     #         AL, caches = part1.l_model_forward(X_batch, parameters, use_batchnorm=True)
#     #
#     #         # Compute cost
#     #         cost = part1.compute_cost(AL, Y_batch)
#     #
#     #
#     #         # Backward propagation
#     #         grads = part2.l_model_backward(AL, Y_batch, caches)
#     #
#     #         # Update parameters
#     #         parameters = part2.update_parameters(parameters, grads, learning_rate)
#     #
#     #         step_counter += 1
#     #
#     #         if step_counter % 100 == 0:
#     #             print(f"Cost after iteration {step_counter}: {cost:.4f}")
#     #             costs.append(cost)
#     #
#     #         if step_counter % 10 == 0: # every 10 mini batches evaluate the model on the validation set
#     #             val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=True)
#     #
#     #             if val_accuracy > max_val_accuracy + small_improvement:
#     #                 max_val_accuracy = val_accuracy
#     #                 no_improve_count = 0
#     #             else:
#     #                 no_improve_count += 1
#     #
#     #             # print(f"Step {step_counter} | Cost: {cost:.4f} | Val Acc: {val_accuracy:.2f}%")
#     #             if no_improve_count > num_steps_without_improvement:
#     #                 print("Early stop no improvement on the validation")
#     #                 break
#     #     if no_improve_count > num_steps_without_improvement:
#     #         break
#     # print("Training finished.")
#     # train_accuracy = part3.predict(X_train, Y_train, parameters, use_batchnorm=True)
#     # val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=True)
#     # test_accuracy = part3.predict(X_test, Y_test, parameters, use_batchnorm=True)
#     #
#     # print(f"Train accuracy: {train_accuracy:.2f}%")
#     # print(f"Validation accuracy: {val_accuracy:.2f}%")
#     # print(f"Test accuracy: {test_accuracy:.2f}%")
#
#
# def train_model_section_6(X_train, Y_train, X_val, Y_val, X_test, Y_test, layers, learning_rate, batch_size, epochs):
#     max_val_accuracy = 0
#     no_improve_count = 0
#     num_steps_without_improvement = 100
#     small_improvement = 1e-4
#     costs = []
#     num_of_iterations = X_train.shape[1] // batch_size
#     parameters = None
#     for epoch in range(epochs):
#         print(f"Epoch {epoch + 1}/{epochs}")
#         parameters, epoch_costs = part3.l_layer_model(X_train, Y_train, layers, learning_rate, num_of_iterations,
#                                                       batch_size, use_batchnorm=False, use_l2=True,lambd=0.01,
#                                                       parameters=parameters)
#         costs.extend((epoch,epoch_costs))
#
#         val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=False)
#         print(f"Validation accuracy after epoch {epoch + 1}: {val_accuracy:.2f}%")
#         if val_accuracy > max_val_accuracy + small_improvement:
#             max_val_accuracy = val_accuracy
#             no_improve_count = 0
#         else:
#             no_improve_count += 1
#
#         if no_improve_count > num_steps_without_improvement:
#             print("Early stop no improvement on the validation")
#             break
#
#     print("Training finished.")
#     train_accuracy = part3.predict(X_train, Y_train, parameters)
#     val_accuracy = part3.predict(X_val, Y_val, parameters)
#     test_accuracy = part3.predict(X_test, Y_test, parameters)
#
#     print(f"Train accuracy: {train_accuracy:.2f}%")
#     print(f"Validation accuracy: {val_accuracy:.2f}%")
#     print(f"Test accuracy: {test_accuracy:.2f}%")
#
#     print("###############################")
#     print("Weights Norms Per Layer section 6:")
#     for l in range(1, len(layers)):
#         W = parameters[f'W{l}']
#         print(f"||W{l}||_F = {np.linalg.norm(W):.4f}")
#
#     return parameters, costs

def section_4():
    X_train, Y_train, X_val, Y_val, X_test, Y_test = get_dataset()
    # Initialize parameters first layer is the input with size 784 (flattened 28x28)
    layers = [784, 20, 7, 5, 10]
    learning_rate = 0.009
    batch_size = 64
    epochs = 500
    parameters_4, costs_4 = train_model_generic(X_train, Y_train, X_val, Y_val, X_test, Y_test,
                                            layers, learning_rate, batch_size, epochs,
                                            use_batchnorm=False, use_l2=False,
                                            print_weights_norms=True)

    write_costs_to_docx(costs_4, "section_4_costs.docx")


def section_5():
    X_train, Y_train, X_val, Y_val, X_test, Y_test = get_dataset()
    # Initialize parameters first layer is the input with size 784 (flattened 28x28)
    layers = [784, 20, 7, 5, 10]
    learning_rate = 0.009
    batch_size = 64
    epochs = 500
    parameters_5, costs_5 = train_model_generic(X_train, Y_train, X_val, Y_val, X_test, Y_test,
                                                layers, learning_rate, batch_size, epochs,
                                                use_batchnorm=True, use_l2=False)
    write_costs_to_docx(costs_5, "section_5_costs.docx")



def section_6():
    X_train, Y_train, X_val, Y_val, X_test, Y_test = get_dataset()
    # Initialize parameters first layer is the input with size 784 (flattened 28x28)
    layers = [784, 20, 7, 5, 10]
    learning_rate = 0.009
    batch_size = 64
    epochs = 500
    parameters_6, costs_6 = train_model_generic(X_train, Y_train, X_val, Y_val, X_test, Y_test,
                                                layers, learning_rate, batch_size, epochs,
                                                use_batchnorm=False, use_l2=True, lambd=0.01,
                                                print_weights_norms=True)
    write_costs_to_docx(costs_6,"section_6_costs.docx")


from docx import Document

def write_costs_to_docx(costs, filename="training_costs.docx"):
    doc = Document()
    doc.add_heading("Training Costs by Epoch and Step", level=1)

    # Create a table with 3 columns: Epoch, Step, Cost
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Epoch"
    hdr_cells[1].text = "Step"
    hdr_cells[2].text = "Cost"

    # Flatten and write the table
    for epoch, epoch_costs in costs:
        for step, cost in epoch_costs:
            row_cells = table.add_row().cells
            row_cells[0].text = str(epoch)
            row_cells[1].text = str(step)
            row_cells[2].text = f"{cost:.4f}"

    doc.save(filename)


def train_model_generic(X_train, Y_train, X_val, Y_val, X_test, Y_test, layers,
                        learning_rate, batch_size, epochs,
                        use_batchnorm=False, use_l2=False, lambd=0.0,
                        print_weights_norms=False):

    max_val_accuracy = 0
    no_improve_count = 0
    num_steps_without_improvement = 100
    small_improvement = 1e-4
    costs = []
    num_of_iterations = X_train.shape[1] // batch_size
    parameters = None

    for epoch in range(epochs):
        print(f"Epoch {epoch + 1}/{epochs}")

        parameters, epoch_costs = part3.l_layer_model(
            X_train, Y_train, layers, learning_rate,
            num_of_iterations, batch_size,
            use_batchnorm=use_batchnorm,
            use_l2=use_l2,
            lambd=lambd,
            parameters=parameters
        )

        costs.append((epoch, epoch_costs))

        # Evaluate on the validation set
        val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=use_batchnorm)
        print(f"Validation accuracy after epoch {epoch + 1}: {val_accuracy:.2f}%")

        # Early stopping logic
        if val_accuracy > max_val_accuracy + small_improvement:
            max_val_accuracy = val_accuracy
            no_improve_count = 0
        else:
            no_improve_count += 1

        if no_improve_count > num_steps_without_improvement:
            print("Early stop: no improvement on the validation set")
            break

    print("Training finished.")
    # Evaluate on all sets
    train_accuracy = part3.predict(X_train, Y_train, parameters, use_batchnorm=use_batchnorm)
    val_accuracy = part3.predict(X_val, Y_val, parameters, use_batchnorm=use_batchnorm)
    test_accuracy = part3.predict(X_test, Y_test, parameters, use_batchnorm=use_batchnorm)

    print(f"Train accuracy: {train_accuracy:.2f}%")
    print(f"Validation accuracy: {val_accuracy:.2f}%")
    print(f"Test accuracy: {test_accuracy:.2f}%")

    if print_weights_norms:
        print("###############################")
        print("Weights Norms Per Layer:")
        for l in range(1, len(layers)):
            W = parameters[f'W{l}']
            print(f"||W{l}||_F = {np.linalg.norm(W):.4f}")

    return parameters, costs


def main():
    print("Section 4 - Without batchnorm")
    print('\n')
    start_time_4 = time.time()
    section_4()
    elapsed_time_4 = time.time() - start_time_4
    print("Elapsed time for section 4: {:.2f} seconds".format(elapsed_time_4))
    print("\nSection 5 - With batchnorm")
    print('\n')
    start_time_5 = time.time()
    section_5()
    elapsed_time_5 = time.time() - start_time_5
    print("Elapsed time for section 5: {:.2f} seconds".format(elapsed_time_5))
    print("\nSection 6 - With L2 regularization")
    print('\n')
    start_time_6 = time.time()
    section_6()
    elapsed_time_6 = time.time() - start_time_6
    print("Elapsed time for section 6: {:.2f} seconds".format(elapsed_time_6))
    print("###############################")
if __name__ == '__main__':
    main()
